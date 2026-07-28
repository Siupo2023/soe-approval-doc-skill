#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
国企请示文档生成器
根据国有企业公文格式规范自动生成 Word 文档
"""

import os
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== 格式常量定义 ==============

# 页边距（毫米转换为 Twips，1mm ≈ 56.7 twips）
MARGIN_TOP_MM = 25
MARGIN_BOTTOM_MM = 25
MARGIN_LEFT_MM = 28
MARGIN_RIGHT_MM = 26

# 字号（磅）
TITLE_FONT_SIZE = 22  # 二号字
BODY_FONT_SIZE = 16   # 三号字

# 行距（磅）
TITLE_LINE_SPACING = 30  # 标题行距
BODY_LINE_SPACING = 28  # 正文行距

# 字体名称（按优先级）
TITLE_FONTS = ['FZXiaoBiaoSong-B05S', '方正小标宋简体', '方正小标宋', '华文中宋', 'SimHei']
BODY_FONTS = ['FangSong_GB2312', 'FZFangSong-Z02', '仿宋GB2312', '仿宋', 'SimSun']

# 段落缩进（字符，约等于 2 字符 * 12 磅/字符 ≈ 24 磅 ≈ 0.85 cm）
FIRST_LINE_INDENT_CHARS = 2


def set_margins(section, top_mm, bottom_mm, left_mm, right_mm):
    """设置页边距（单位：毫米）"""
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(top_mm)
    section.bottom_margin = Mm(bottom_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)


def set_chinese_font(run, font_names, size_pt, bold=False):
    """设置中文字体（支持备选字体列表）"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold

    # 尝试设置中文字体
    for font_name in font_names:
        try:
            run.font.name = font_name
            # 设置东亚字体（对中文很重要）
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            break
        except:
            continue


def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        line_spacing_pt=None, first_line_indent_chars=None,
                        space_before=None, space_after=None):
    """设置段落格式"""
    # 对齐方式
    paragraph.alignment = alignment

    # 行距
    if line_spacing_pt:
        paragraph.paragraph_format.line_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(line_spacing_pt)

    # 首行缩进
    if first_line_indent_chars is not None:
        # 2 字符缩进（约等于 0.85 cm 或 24 磅）
        paragraph.paragraph_format.first_line_indent = (
            Inches(0.33) if first_line_indent_chars else Inches(0)
        )

    # 段前段后间距
    if space_before is not None:
        paragraph.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)


def add_title(doc, text):
    """添加标题（方正小标宋简体，二号字，30磅行距，居中）"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_chinese_font(run, TITLE_FONTS, TITLE_FONT_SIZE)
    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       line_spacing_pt=TITLE_LINE_SPACING)
    return paragraph


def add_salutation(doc, text):
    """添加称谓（如"公司领导："）"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       line_spacing_pt=BODY_LINE_SPACING)
    return paragraph


def add_body_paragraph(doc, text):
    """添加正文段落（仿宋GB2312，三号字，28磅行距，首行缩进2字符）"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       line_spacing_pt=BODY_LINE_SPACING,
                       first_line_indent_chars=FIRST_LINE_INDENT_CHARS)
    return paragraph


def add_list_paragraph(doc, text, numbered=False):
    """添加使用 Word 原生列表样式的条目。"""
    style = 'List Number' if numbered else 'List Bullet'
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing_pt=BODY_LINE_SPACING,
        first_line_indent_chars=0,
    )
    return paragraph


def add_labeled_paragraph(doc, label, text):
    """添加“标签：内容”段落，标签加粗且不保留 Markdown 标记。"""
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}：")
    set_chinese_font(label_run, BODY_FONTS, BODY_FONT_SIZE, bold=True)
    text_run = paragraph.add_run(text)
    set_chinese_font(text_run, BODY_FONTS, BODY_FONT_SIZE)
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing_pt=BODY_LINE_SPACING,
        first_line_indent_chars=0,
    )
    return paragraph


def chinese_number(number):
    """将 1-99 转为公文标题常用中文数字。"""
    digits = '零一二三四五六七八九'
    if number < 10:
        return digits[number]
    tens, ones = divmod(number, 10)
    prefix = '' if tens == 1 else digits[tens]
    return f"{prefix}十{digits[ones] if ones else ''}"


def add_section_heading(doc, level, text, number=1):
    """
    添加小标题

    level: 1=一、, 2=（一）, 3=1., 4=（1）
    """
    if level < 1 or level > 4:
        level = 1

    prefixes = {
        1: f"{chinese_number(number)}、",
        2: f"（{chinese_number(number)}）",
        3: f"{number}.",
        4: f"（{number}）",
    }
    full_text = f"{prefixes[level]}{text}"

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(full_text)
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE, bold=(level == 1))

    # 一级标题不缩进，其他级别缩进
    indent = FIRST_LINE_INDENT_CHARS if level > 1 else None

    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       line_spacing_pt=BODY_LINE_SPACING,
                       first_line_indent_chars=indent)
    return paragraph


def add_request_items(doc, items):
    """添加请示事项（不缩进的列表）"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("现将有关事项请示如下：")
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       line_spacing_pt=BODY_LINE_SPACING)

    for i, item in enumerate(items, 1):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{i}、{item}")
        set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
        set_paragraph_format(paragraph,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           line_spacing_pt=BODY_LINE_SPACING,
                           first_line_indent_chars=FIRST_LINE_INDENT_CHARS)

    # 结束语
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("妥否，请批示。")
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       line_spacing_pt=BODY_LINE_SPACING,
                       first_line_indent_chars=FIRST_LINE_INDENT_CHARS)


def add_attachments(doc, attachments):
    """
    添加附件说明

    attachments: list of attachment names
    格式：正文下空 1 行，左空 2 字
    """
    # 添加空行
    doc.add_paragraph()

    paragraph = doc.add_paragraph()
    text = "附件：" + " ".join([f"{i+1}.{name}" for i, name in enumerate(attachments)])
    run = paragraph.add_run(text)
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)

    # 左空 2 字（约等于 0.85 cm 或 24 磅）
    paragraph.paragraph_format.left_indent = Inches(0.33)
    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       line_spacing_pt=BODY_LINE_SPACING)


def add_signature(doc, department, date_str=None):
    """
    添加落款

    department: 部门名称
    date_str: 日期字符串（格式："YYYY年MM月DD日"），默认为当天
    """
    if date_str is None:
        date_str = datetime.now().strftime("%Y年%m月%d日")

    # 添加一些空行
    for _ in range(2):
        doc.add_paragraph()

    # 部门名称
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(department)
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                       line_spacing_pt=BODY_LINE_SPACING)

    # 日期
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(date_str)
    set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
    set_paragraph_format(paragraph,
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                       line_spacing_pt=BODY_LINE_SPACING)


def generate_approval_doc(output_path, title, salutation, sections,
                         request_items, attachments, department):
    """
    生成完整的请示文档

    参数:
        output_path: 输出文件路径（.docx）
        title: 文档标题
        salutation: 称谓（如"公司领导："）
        sections: list of dict，每个 dict 包含:
            - level: 标题层级 (1-4)
            - heading: 标题文本
            - content: list of 段落文本
        request_items: list of 请示事项
        attachments: list of 附件名称
        department: 落款部门
    """
    doc = Document()

    # 设置页边距
    set_margins(doc.sections[0],
               MARGIN_TOP_MM,
               MARGIN_BOTTOM_MM,
               MARGIN_LEFT_MM,
               MARGIN_RIGHT_MM)

    # 标题
    add_title(doc, title)

    # 称谓
    add_salutation(doc, salutation)

    # 各个部分
    heading_counters = {1: 0, 2: 0, 3: 0, 4: 0}
    for section in sections:
        level = section.get('level', 1)
        heading = section.get('heading', '')
        content = section.get('content', [])
        heading_counters[level] += 1
        for deeper_level in range(level + 1, 5):
            heading_counters[deeper_level] = 0

        # 小标题
        add_section_heading(doc, level, heading, heading_counters[level])

        # 正文段落
        for para_text in content:
            if para_text.strip():
                add_body_paragraph(doc, para_text)

    # 请示事项
    if request_items:
        add_request_items(doc, request_items)

    # 附件
    if attachments:
        add_attachments(doc, attachments)

    # 落款
    add_signature(doc, department)

    # 保存文档
    output_file = Path(output_path).expanduser()
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)
    print(f"✅ 文档已生成：{output_file}")


# ============== 工作总结生成器（新增）=============

def generate_work_summary_doc(output_path,
                               title,
                               department_info,
                               overview,
                               achievements,
                               problems,
                               experiences,
                               plans,
                               signature):
    """
    生成工作总结文档

    参数:
        output_path: 输出文件路径（.docx）
        title: 文档标题（如"2025年度工作总结"）
        department_info: dict，部门信息
            - department: 部门/岗位
            - date_range: 时间范围
        overview: str，工作概述
        achievements: list of dict，主要成绩
            - area: 工作领域
            - items: list of 成果条目
        problems: list of dict，存在问题
            - title: 问题标题
            - details: list of 问题细节
        experiences: list of str，经验总结
        plans: list of dict，下一步计划
            - goal: 目标
            - measures: 措施
            - timeline: 时间节点
        signature: dict，落款信息
            - name: 姓名/部门
            - date: 日期
    """
    doc = Document()

    # 设置页边距
    set_margins(doc.sections[0],
               MARGIN_TOP_MM,
               MARGIN_BOTTOM_MM,
               MARGIN_LEFT_MM,
               MARGIN_RIGHT_MM)

    # 标题
    add_title(doc, title)

    # 基本信息（部门/岗位、时间）
    if department_info:
        department = department_info.get('department', '')
        name = department_info.get('name', '')
        identity = '：'.join(part for part in (department, name) if part)
        for info_text in (identity, department_info.get('date_range', '')):
            if not info_text:
                continue
            para = doc.add_paragraph()
            run = para.add_run(info_text)
            set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
            set_paragraph_format(
                para,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                line_spacing_pt=BODY_LINE_SPACING,
            )

    # 一、工作概述
    if overview:
        add_section_heading(doc, 1, '工作概述', 1)
        add_body_paragraph(doc, overview)

    # 二、主要成绩
    if achievements:
        add_section_heading(doc, 1, '主要成绩', 2)
        for ach_index, ach in enumerate(achievements, 1):
            area = ach.get('area', '')
            items = ach.get('items', [])

            # 工作领域标题
            add_section_heading(doc, 2, area, ach_index)

            # 成果列表
            for item in items:
                add_list_paragraph(doc, item)

    # 三、存在问题
    if problems:
        add_section_heading(doc, 1, '存在问题', 3)
        for problem_index, prob in enumerate(problems, 1):
            prob_title = prob.get('title', '')
            details = prob.get('details', [])

            # 问题标题
            add_section_heading(doc, 2, prob_title, problem_index)

            # 问题细节
            for detail in details:
                add_labeled_paragraph(
                    doc,
                    detail.get("label", ""),
                    detail.get("text", ""),
                )

    # 四、经验总结
    if experiences:
        add_section_heading(doc, 1, '经验总结', 4)
        for exp in experiences:
            add_list_paragraph(doc, exp, numbered=True)

    # 五、下一步计划
    if plans:
        add_section_heading(doc, 1, '下一步计划', 5)
        for plan_index, plan in enumerate(plans, 1):
            goal = plan.get('goal', '')
            measures = plan.get('measures', '')
            timeline = plan.get('timeline', '')

            # 计划标题
            add_section_heading(doc, 2, goal, plan_index)

            # 计划细节
            add_labeled_paragraph(doc, '目标', goal)
            add_labeled_paragraph(doc, '措施', measures)
            add_labeled_paragraph(doc, '时间节点', timeline)

    # 落款
    if signature:
        for _ in range(3):
            doc.add_paragraph()

        # 姓名/部门
        para = doc.add_paragraph()
        run = para.add_run(signature.get('name', ''))
        set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
        set_paragraph_format(para,
                           alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                           line_spacing_pt=BODY_LINE_SPACING)

        # 日期
        para = doc.add_paragraph()
        run = para.add_run(signature.get('date', datetime.now().strftime("%Y年%m月%d日")))
        set_chinese_font(run, BODY_FONTS, BODY_FONT_SIZE)
        set_paragraph_format(para,
                           alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                           line_spacing_pt=BODY_LINE_SPACING)

    # 保存文档
    output_file = Path(output_path).expanduser()
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)
    print(f"✅ 工作总结已生成：{output_file}")


# ============== 命令行接口 ==============

def main():
    """命令行接口"""
    if len(sys.argv) < 2:
        print("用法: python3 generate_approval_doc.py <输出文件.docx>")
        print("或从 Claude 调用 generate_approval_doc() 函数")
        sys.exit(1)

    output_path = sys.argv[1]

    # 示例：生成一个测试文档
    generate_approval_doc(
        output_path=output_path,
        title="关于采购办公设备的请示",
        salutation="公司领导：",
        sections=[
            {
                'level': 1,
                'heading': '基本情况',
                'content': [
                    '我部门现有办公设备已使用超过5年，部分设备已出现老化现象，影响日常办公效率。',
                    '为保障部门正常运转，提高工作效率，现申请更新部分办公设备。'
                ]
            },
            {
                'level': 1,
                'heading': '必要性与可行性',
                'content': [
                    '必要性：现有设备故障频发，维修成本逐年增加，影响工作进度。',
                    '可行性：经市场调研，相关设备价格合理，预算在部门年度预算范围内。'
                ]
            },
            {
                'level': 1,
                'heading': '采购方案',
                'content': [
                    '拟采购电脑10台，打印机2台，投影仪1台。',
                    '预算总金额约15万元。'
                ]
            }
        ],
        request_items=[
            '同意采购办公设备，预算金额15万元',
            '同意按规定程序进行采购'
        ],
        attachments=['办公设备采购清单', '报价单'],
        department='综合管理部'
    )


if __name__ == '__main__':
    main()
